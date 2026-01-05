import sys
import os
import json
import sqlite3
import requests
import urllib3
import time
import numpy as np
import re
import hashlib
import csv

# ================= 新增导出所需的库 =================
# 为了保证代码健壮性，尝试导入第三方导出库，如果缺失则在运行时提示
try:
    import pandas as pd
except ImportError:
    pd = None

try:
    import docx
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    docx = None

from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, 
                             QLabel, QLineEdit, QPushButton, QTextEdit, QFileDialog, 
                             QMessageBox, QSplitter, QFrame, QComboBox, QProgressBar)
from PyQt5.QtCore import QThread, pyqtSignal, Qt, QSettings
from PyQt5.QtGui import QTextCursor

# ================= 配置与环境 =================
# 禁用 HTTPS 警告 (Win7/内网适配)
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
os.environ['CURL_CA_BUNDLE'] = ''

# API 配置 (硬编码 Key)
API_KEY = "your api key"

# 1. Embedding API
EMBEDDING_API_URL = "https://www.bge.com/v1/embeddings" 
EMBEDDING_MODEL_NAME = "bge-m3"

# 2. Rerank API
RERANK_API_URL = "https://www.bge.com/v1/rerank" 
RERANK_MODEL_NAME = "bge-reranker-v2-m3"

# 3. DeepSeek R1 API
DEEPSEEK_API_URL = "https://www.deepseek.com/v1/chat/completions"
DEEPSEEK_MODEL_NAME = "DeepSeek-R1"

# ================= System Prompt (DeepSeek) =================
DEEPSEEK_SYSTEM_PROMPT = """🎯【角色定义】
你是一个 RAG Final Answer Composer（检索增强生成的最终答案生成器）。
你的任务 不是检索、不是排序、不是猜测，而是：
严格基于已提供的召回结果，对用户问题生成最终、可读、准确的回答。

📥【输入说明】
你将收到一个结构化输入，包含：
1. query: 用户的原始问题
2. retrieved_chunks（Top-K，已完成向量召回 + reranker 排序）

🔒【强制约束（非常重要）】
1️⃣ 事实来源约束（防幻觉）
❌ 禁止 使用任何外部知识
❌ 禁止 补充未在 retrieved_chunks 中出现的事实
✅ 只允许 基于提供内容进行归纳、重写、总结
如果证据不足：必须明确说明「当前召回内容不足以完整回答该问题」

2️⃣ 内容使用规则（防遗漏）
优先使用 Rank 靠前的内容
若多个 chunk 语义重复，应：合并信息、去除重复表述
不得忽略与 query 明确相关的高分 chunk

3️⃣ 噪声处理规则（适配 PDF / OCR）
允许你：修复断行、合并被拆散的句子、去除明显乱码
❌ 不允许“合理猜测”缺失内容

✍️【输出要求】
输出必须满足：
✅ 语言清晰、技术准确
✅ **必须使用 Markdown 格式，包含清晰的段落、列表和加粗**
✅ 不直接大段复制原文（允许短引用）
✅ 不提及“召回 / reranker / 向量 / chunk”等系统概念

📐【推荐输出结构（自动选择）】
根据问题复杂度，自适应选择：
- 简单问题：直接给出 1–2 段 concise 回答
- 技术型问题（推荐）：简要结论（1–2 句） + 详细说明（要点列表） + 补充说明

⚠️【失败兜底策略】
如果所有 retrieved_chunks 与 query 相关性都很弱，或内容彼此矛盾、无法整合，
你必须输出：“根据当前召回的文档内容，无法对该问题给出可靠回答。”

✅【总结一句话】
你是一个“只基于证据的答案生成器”，不是一个自由发挥的聊天模型。"""

# ================= 样式表 (Dark Mode) =================
STYLESHEET = """
QMainWindow { background-color: #2b2b2b; color: #e0e0e0; font-family: 'Segoe UI', 'Microsoft YaHei', sans-serif; }
QLabel { color: #aaaaaa; font-weight: bold; font-size: 13px; }
QLineEdit { background-color: #3c3c3c; color: #ffffff; border: 1px solid #555555; padding: 6px; border-radius: 4px; }
QTextEdit { background-color: #1e1e1e; color: #e0e0e0; border: 1px solid #444444; font-family: Consolas, monospace; font-size: 12px; }
QPushButton { background-color: #007acc; color: white; border: none; padding: 8px 16px; border-radius: 4px; font-weight: bold; font-size: 14px; }
QPushButton:hover { background-color: #005f9e; }
QPushButton:pressed { background-color: #004a80; }
QPushButton:disabled { background-color: #444444; color: #888888; }
QComboBox { background-color: #3c3c3c; color: white; border: 1px solid #555; padding: 5px; border-radius: 4px; }
QComboBox::drop-down { border: 0px; }
QFrame#Divider { border: 1px solid #444444; }
"""

# ================= 核心工具：相似度计算 =================
def cosine_similarity(vec1, vec2):
    try:
        norm1 = np.linalg.norm(vec1)
        norm2 = np.linalg.norm(vec2)
        if norm1 == 0 or norm2 == 0:
            return 0.0
        return np.dot(vec1, vec2) / (norm1 * norm2)
    except Exception:
        return 0.0

# ================= 核心工具：内容哈希去重 =================
def get_text_hash(text):
    """生成文本的 SHA256 哈希，用于严格去重"""
    return hashlib.sha256(text.encode('utf-8')).hexdigest()

# ================= PageIndex Loader =================
class PageIndexLoader:
    def __init__(self):
        self.index = {}          
        self.ordered_ids = []    
        self.is_loaded = False

    def load_json(self, json_path):
        if not json_path or not os.path.exists(json_path):
            return False, "文件不存在"
        
        try:
            with open(json_path, 'r', encoding='utf-8-sig') as f:
                data = json.load(f)
            
            self.index = {}
            self.ordered_ids = []
            
            root_structure = data.get("structure", []) if isinstance(data, dict) else data
            
            for item in root_structure:
                self._traverse(item, parent_path=[])
            
            self.is_loaded = True
            return True, f"成功加载 PageIndex，包含 {len(self.index)} 个节点"
        except Exception as e:
            return False, f"加载异常: {str(e)}"

    def _traverse(self, node, parent_path):
        current_title = node.get("title", "")
        node_id = str(node.get("node_id", "")) 
        current_path = parent_path + [current_title]
        
        if node_id:
            self.index[node_id] = {
                "title": current_title,
                "text": node.get("text", ""),
                "summary": node.get("summary", ""),
                "path": current_path,
                "raw_node": node 
            }
            self.ordered_ids.append(node_id)
        
        if "nodes" in node and isinstance(node["nodes"], list):
            for child in node["nodes"]:
                self._traverse(child, current_path)

    def get_node(self, node_id):
        return self.index.get(str(node_id))

# ================= 工作线程：工业级鲁棒召回 + DeepSeek 总结 =================
class RecallWorker(QThread):
    log_signal = pyqtSignal(str)          
    result_signal = pyqtSignal(list)      
    summary_signal = pyqtSignal(str)      # 发送 DeepSeek 总结内容
    finish_signal = pyqtSignal(bool)      

    def __init__(self, query_text, db_path, json_path):
        super().__init__()
        self.query_text = query_text
        self.db_path = db_path
        self.json_path = json_path
        self.page_index = PageIndexLoader()

    def log(self, msg):
        timestamp = time.strftime("%H:%M:%S")
        self.log_signal.emit(f"[{timestamp}] {msg}")

    # --- Step 1: Embedding ---
    def get_remote_embedding(self, text):
        self.log(f"正在发送 Query 到 BGE-M3: {text[:20]}...")
        headers = { 'Content-Type': 'application/json', 'Authorization': f'Bearer {API_KEY}' }
        payload = { "model": EMBEDDING_MODEL_NAME, "input": [text] }
        
        try:
            response = requests.post(EMBEDDING_API_URL, headers=headers, json=payload, verify=False, timeout=30)
            if response.status_code == 200:
                data = response.json()
                if 'data' in data and len(data['data']) > 0:
                    return data['data'][0]['embedding']
        except Exception as e:
            self.log(f"❌ Embedding 网络异常: {str(e)}")
        return None

    # --- Step 2: Rerank API ---
    def rerank_with_bge(self, query, candidates_text_list):
        self.log(f"📡 Reranker ({RERANK_MODEL_NAME}) 正在处理 {len(candidates_text_list)} 条数据...")
        headers = { 'Content-Type': 'application/json', 'Authorization': f'Bearer {API_KEY}' }
        
        payload = {
            "model": RERANK_MODEL_NAME,
            "query": query,
            "documents": candidates_text_list 
        }

        try:
            start_time = time.time()
            response = requests.post(RERANK_API_URL, headers=headers, json=payload, verify=False, timeout=30)
            
            if response.status_code == 200:
                data = response.json()
                scores = [0.0] * len(candidates_text_list)
                
                if "results" in data:
                    for res in data["results"]:
                        idx = res.get("index")
                        score = res.get("relevance_score", 0.0)
                        if idx is not None and 0 <= idx < len(scores):
                            scores[idx] = score
                elif isinstance(data, list):
                     scores = data
                else:
                    self.log("⚠️ Reranker 返回格式未知，降级处理。")
                    return None

                self.log(f"✅ Reranker 完成，耗时: {time.time() - start_time:.2f}s")
                return scores
            else:
                self.log(f"⚠️ Reranker 请求失败: {response.status_code}")
                return None
        except Exception as e:
            self.log(f"⚠️ Reranker 调用异常: {str(e)}")
            return None

    # --- Step 3: 规则裁决 ---
    def apply_industrial_rules(self, query, path_str, original_score):
        q_lower = query.lower()
        p_lower = path_str.lower()
        final_adj_score = original_score

        technical_terms = ["train", "optimi", "loss", "layer", "struct", "arch"]
        if any(t in q_lower for t in technical_terms):
            negative_sections = ["introduction", "background", "preface", "motivation", "overview", "why", "related work"]
            for neg in negative_sections:
                if neg in p_lower:
                    final_adj_score -= 3.0 
                    break
            if "train" in q_lower and "train" in p_lower:
                final_adj_score += 1.0
        
        return final_adj_score

    # --- Step 4: DeepSeek R1 Summary (流式) ---
    def call_deepseek_summary(self, query, top_results):
        self.log("🧠 正在请求 DeepSeek-R1 生成总结 (Stream=True)...")
        self.summary_signal.emit("> 🚀 **DeepSeek-R1 已连接，准备生成...**\n\n")

        # 1. 构造 Context
        context_str = ""
        for item in top_results:
            context_str += f"""
---
[Rank {item['rank']}] (Score: {item['final_score']:.2f})
Section Path: {item['path']}
Content:
{item['content']}
"""
        
        user_prompt_content = f"Query: {query}\n\nRetrieved Chunks:{context_str}"

        headers = { 'Content-Type': 'application/json', 'Authorization': f'Bearer {API_KEY}' }
        payload = {
            "model": DEEPSEEK_MODEL_NAME,
            "messages": [
                {"role": "system", "content": DEEPSEEK_SYSTEM_PROMPT},
                {"role": "user", "content": user_prompt_content}
            ],
            "stream": True,
            "temperature": 0.6
        }

        try:
            response = requests.post(DEEPSEEK_API_URL, headers=headers, json=payload, verify=False, stream=True)
            
            if response.status_code == 200:
                full_reasoning = ""
                full_content = ""
                
                for line in response.iter_lines():
                    if line:
                        decoded_line = line.decode('utf-8')
                        if decoded_line.startswith("data: "):
                            data_str = decoded_line[6:] 
                            if data_str.strip() == "[DONE]":
                                break
                            try:
                                json_chunk = json.loads(data_str)
                                delta = json_chunk['choices'][0]['delta']
                                
                                current_reasoning_delta = delta.get('reasoning_content', '')
                                current_content_delta = delta.get('content', '')
                                updated = False
                                
                                if current_reasoning_delta:
                                    full_reasoning += current_reasoning_delta
                                    updated = True
                                if current_content_delta:
                                    full_content += current_content_delta
                                    updated = True

                                if updated:
                                    formatted_output = ""
                                    if full_reasoning:
                                        clean_reasoning = full_reasoning.replace('\n', '\n> ')
                                        formatted_output += f"> 🧠 **DeepSeek Thinking Process:**\n> {clean_reasoning}\n\n"
                                    
                                    if full_content:
                                        if full_reasoning:
                                            formatted_output += "---\n\n" 
                                        formatted_output += f"{full_content}"
                                        
                                    self.summary_signal.emit(formatted_output)
                            except Exception:
                                continue
                self.log("✅ DeepSeek 总结生成完毕")
            else:
                self.log(f"❌ DeepSeek API 错误: {response.status_code}")
                self.summary_signal.emit(f"⚠️ 无法生成总结: API Error {response.status_code}")

        except Exception as e:
            self.log(f"❌ DeepSeek 调用异常: {str(e)}")
            self.summary_signal.emit(f"⚠️ 总结生成失败: {str(e)}")


    def run(self):
        try:
            # 0. 加载 PageIndex
            has_pageindex = False
            if self.json_path:
                self.log(f"加载 PageIndex: {os.path.basename(self.json_path)}...")
                success, msg = self.page_index.load_json(self.json_path)
                if success:
                    has_pageindex = True
                else:
                    self.log(f"⚠️ PageIndex 加载失败: {msg}")

            # 1. Query Vector
            query_vec_list = self.get_remote_embedding(self.query_text)
            if not query_vec_list:
                self.finish_signal.emit(False)
                return
            query_vec_np = np.array(query_vec_list, dtype=np.float32)

            # 2. SQLite Vector Search
            if not os.path.exists(self.db_path):
                self.log("❌ 数据库不存在")
                self.finish_signal.emit(False)
                return

            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            cursor.execute("SELECT id, embedding, section_id FROM vectors")
            rows = cursor.fetchall()
            
            raw_candidates = []
            for row in rows:
                doc_id, emb_json, sec_id_db = row
                try:
                    doc_vec = np.array(json.loads(emb_json), dtype=np.float32)
                    score = cosine_similarity(query_vec_np, doc_vec)
                    raw_candidates.append({"id": doc_id, "vec_score": score, "section_id": str(sec_id_db)})
                except: continue

            raw_candidates.sort(key=lambda x: x["vec_score"], reverse=True)
            top_candidates_raw = raw_candidates[:30] 

            # 3. 去重 & 构造输入
            rerank_input_texts = [] 
            processed_candidates = [] 
            seen_content_hashes = set()

            for item in top_candidates_raw:
                sec_id = item["section_id"]
                node_info = self.page_index.get_node(sec_id) if has_pageindex else None
                
                raw_text = ""
                path_str = ""
                summary_text = ""

                if node_info:
                    raw_text = node_info['text']
                    path_str = " > ".join(node_info['path'])
                    summary_text = node_info.get('summary', '')
                else:
                    cursor.execute("SELECT original_snippet, section_path FROM documents WHERE id=?", (item['id'],))
                    db_row = cursor.fetchone()
                    if db_row:
                        raw_text = db_row[0]
                        path_str = str(db_row[1])

                content_hash = get_text_hash(raw_text)
                if content_hash in seen_content_hashes:
                    continue 
                seen_content_hashes.add(content_hash)

                context_aware_input = f"Section Path: {path_str}\nContent: {raw_text}"
                rerank_input_texts.append(context_aware_input)

                display_content = f"[Summary]\n{summary_text}\n\n[Text]\n{raw_text}" if summary_text else raw_text
                
                processed_candidates.append({
                    "id": item["id"],
                    "vec_score": item["vec_score"],
                    "path": path_str,
                    "content": display_content,
                    "final_score": 0.0 
                })
                
                if len(processed_candidates) >= 15:
                    break

            conn.close()

            # 4. 执行 Rerank
            rerank_scores = self.rerank_with_bge(self.query_text, rerank_input_texts)
            
            # 5. 分数融合
            if rerank_scores and len(rerank_scores) == len(processed_candidates):
                for idx, candidate in enumerate(processed_candidates):
                    raw_rerank_score = rerank_scores[idx]
                    adjusted_rerank_score = self.apply_industrial_rules(
                        self.query_text, 
                        candidate['path'], 
                        raw_rerank_score
                    )
                    candidate['final_score'] = 0.2 * candidate['vec_score'] + 0.8 * adjusted_rerank_score
                    candidate['debug_score'] = f"R:{adjusted_rerank_score:.2f} (Orig:{raw_rerank_score:.2f})"
                
                processed_candidates.sort(key=lambda x: x["final_score"], reverse=True)
            else:
                self.log("⚠️ 降级：仅使用向量分排序")
                for candidate in processed_candidates:
                    candidate['final_score'] = candidate['vec_score']
                    candidate['debug_score'] = "VecOnly"

            # 6. Top-10 Result
            final_top_10 = processed_candidates[:10]
            for idx, res in enumerate(final_top_10):
                res['rank'] = idx + 1

            self.result_signal.emit(final_top_10)
            self.call_deepseek_summary(self.query_text, final_top_10)
            self.finish_signal.emit(True)

        except Exception as e:
            self.log(f"❌ 严重错误: {str(e)}")
            import traceback
            self.log(traceback.format_exc())
            self.finish_signal.emit(False)

# ================= 主界面 =================
class RAGRecallApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("RAG 工业级全流程 (Recall -> Rerank -> DeepSeek R1)")
        self.resize(1300, 950) # 略微增加高度适配底部按钮
        self.setStyleSheet(STYLESHEET)
        
        self.settings = QSettings("MyCorp", "RAGRecall_Final_v2")
        # 缓存数据用于导出
        self.cached_results = []
        self.cached_summary = ""
        self.cached_query = ""
        
        self.init_ui()
        
    def init_ui(self):
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        main_layout = QHBoxLayout(central_widget)
        main_layout.setSpacing(10)
        
        # === Left Widget ===
        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)
        
        # DB & JSON Inputs
        db_layout = QHBoxLayout()
        self.db_path_edit = QLineEdit()
        self.db_path_edit.setText(self.settings.value("last_db_path", ""))
        btn_db = QPushButton("📂 向量库")
        btn_db.clicked.connect(self.browse_db)
        db_layout.addWidget(QLabel("Vector DB:"))
        db_layout.addWidget(self.db_path_edit)
        db_layout.addWidget(btn_db)
        left_layout.addLayout(db_layout)
        
        json_layout = QHBoxLayout()
        self.json_path_edit = QLineEdit()
        self.json_path_edit.setText(self.settings.value("last_json_path", ""))
        btn_json = QPushButton("📄 PageIndex")
        btn_json.clicked.connect(self.browse_json)
        json_layout.addWidget(QLabel("Structure JSON:"))
        json_layout.addWidget(self.json_path_edit)
        json_layout.addWidget(btn_json)
        left_layout.addLayout(json_layout)
        
        # Query
        left_layout.addWidget(QLabel("用户查询 (Query):"))
        self.query_input = QTextEdit()
        self.query_input.setPlaceholderText("请输入问题...")
        self.query_input.setMaximumHeight(60)
        left_layout.addWidget(self.query_input)
        
        # Search Button
        self.btn_search = QPushButton("🚀 执行全流程 (Recall + DeepSeek)")
        self.btn_search.setFixedHeight(45)
        self.btn_search.setStyleSheet("background-color: #2da44e; font-size: 15px;")
        self.btn_search.clicked.connect(self.start_recall)
        left_layout.addWidget(self.btn_search)
        
        # DeepSeek Summary
        left_layout.addWidget(QLabel("🤖 DeepSeek-R1 智能总结 (Thinking + Answer):"))
        self.summary_display = QTextEdit()
        self.summary_display.setReadOnly(True)
        self.summary_display.setStyleSheet("""
            QTextEdit {
                background-color: #252526; 
                color: #dcdcaa; 
                font-family: 'Segoe UI', sans-serif; 
                font-size: 14px; 
                border: 1px solid #007acc;
                line-height: 1.6;
            }
        """)
        self.summary_display.setMinimumHeight(300)
        left_layout.addWidget(self.summary_display)

        # Context List
        left_layout.addWidget(QLabel("📚 Reranked Context (Top-10):"))
        self.result_display = QTextEdit()
        self.result_display.setReadOnly(True)
        self.result_display.setStyleSheet("font-family: Consolas; font-size: 12px; color: #aaddff;")
        left_layout.addWidget(self.result_display)
        
        # === 导出功能区域 ===
        export_layout = QHBoxLayout()
        export_layout.addWidget(QLabel("导出格式:"))
        
        self.combo_format = QComboBox()
        self.combo_format.addItems(["xlsx", "csv", "txt", "docx", "md"])
        self.combo_format.setFixedWidth(100)
        export_layout.addWidget(self.combo_format)
        
        self.btn_export = QPushButton("💾 导出结果")
        self.btn_export.setStyleSheet("background-color: #d2691e;")
        self.btn_export.clicked.connect(self.export_data)
        export_layout.addWidget(self.btn_export)
        
        export_layout.addStretch() # 弹簧占位
        left_layout.addLayout(export_layout)

        # === Right Console ===
        right_widget = QWidget()
        right_layout = QVBoxLayout(right_widget)
        right_layout.addWidget(QLabel("📟 System Console"))
        self.console_output = QTextEdit()
        self.console_output.setReadOnly(True)
        self.console_output.setStyleSheet("background-color: #111; color: #0f0; font-family: Consolas;")
        right_layout.addWidget(self.console_output)
        
        splitter = QSplitter(Qt.Horizontal)
        splitter.addWidget(left_widget)
        splitter.addWidget(right_widget)
        splitter.setSizes([900, 400])
        main_layout.addWidget(splitter)

    def browse_db(self):
        path, _ = QFileDialog.getOpenFileName(self, "选择数据库", "", "SQLite DB (*.db);;All Files (*.*)")
        if path:
            self.db_path_edit.setText(path)
            self.settings.setValue("last_db_path", path)

    def browse_json(self):
        path, _ = QFileDialog.getOpenFileName(self, "选择 JSON", "", "JSON Files (*.json);;All Files (*.*)")
        if path:
            self.json_path_edit.setText(path)
            self.settings.setValue("last_json_path", path)

    def log(self, msg):
        self.console_output.append(msg)
        cursor = self.console_output.textCursor()
        cursor.movePosition(QTextCursor.End)
        self.console_output.setTextCursor(cursor)

    def update_summary(self, text):
        """流式更新总结文本"""
        self.cached_summary = text # 实时保存，用于导出
        self.summary_display.setMarkdown(text) 
        cursor = self.summary_display.textCursor()
        cursor.movePosition(QTextCursor.End)
        self.summary_display.setTextCursor(cursor)

    def display_results(self, results):
        self.cached_results = results # 保存结果用于导出
        html = ""
        for item in results:
            score_text = f"{item['final_score']:.4f}"
            debug_info = item.get('debug_score', '')
            color = "#00ff00" if item['final_score'] > 0 else "#ffaa00"
            
            html += f"""
            <div style='border-bottom: 1px solid #555; padding: 12px; margin-bottom: 8px;'>
                <span style='color: #888; font-weight:bold;'>Rank #{item['rank']}</span> | 
                <span style='color: {color}; font-weight: bold;'>Final: {score_text}</span> 
                <span style='color: #aaa; font-size:11px;'>[{debug_info}]</span><br>
                <div style='margin-top:5px; color: #ffcc00;'><b>[Section Path]</b> {item['path']}</div>
                <div style='margin-top:5px; background-color: #222; padding: 8px; border-left: 3px solid #2da44e; white-space: pre-wrap;'>
{item['content'][:200]}...
                </div>
            </div>
            """
        self.result_display.setHtml(html)

    def start_recall(self):
        db_path = self.db_path_edit.text().strip()
        json_path = self.json_path_edit.text().strip()
        query = self.query_input.toPlainText().strip()
        
        if not db_path or not os.path.exists(db_path):
            QMessageBox.warning(self, "Error", "无效的数据库路径")
            return
        if not query:
            return
        
        self.cached_query = query
        self.cached_results = []
        self.cached_summary = ""
        
        self.btn_search.setEnabled(False)
        self.result_display.clear()
        self.summary_display.clear() 
        self.console_output.clear()
        self.log("🚀 初始化任务...")
        
        self.worker = RecallWorker(query, db_path, json_path)
        self.worker.log_signal.connect(self.log)
        self.worker.result_signal.connect(self.display_results)
        self.worker.summary_signal.connect(self.update_summary) 
        self.worker.finish_signal.connect(self.on_finished)
        self.worker.start()

    def on_finished(self, success):
        self.btn_search.setEnabled(True)
        if success:
            self.log("✅ 全流程结束")
        else:
            self.log("❌ 流程失败")

    # ================= 核心导出逻辑 =================
    def export_data(self):
        if not self.cached_summary and not self.cached_results:
            QMessageBox.warning(self, "提示", "当前没有可导出的结果，请先执行查询。")
            return

        fmt = self.combo_format.currentText()
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        filename = f"export_{timestamp}.{fmt}"
        save_path = os.path.join(os.getcwd(), filename)

        try:
            self.log(f"💾 正在导出为 {fmt} ...")
            
            # 1. 导出 XLSX
            if fmt == "xlsx":
                if pd is None:
                    raise ImportError("缺少 pandas 或 openpyxl 库，请 pip install pandas openpyxl")
                
                # 构建数据
                data_rows = []
                for item in self.cached_results:
                    data_rows.append({
                        "Rank": item['rank'],
                        "Final Score": item['final_score'],
                        "Debug Score": item.get('debug_score', ''),
                        "Section Path": item['path'],
                        "Content": item['content']
                    })
                
                df = pd.DataFrame(data_rows)
                
                # 使用 Pandas ExcelWriter 写入
                with pd.ExcelWriter(save_path, engine='openpyxl') as writer:
                    # 将总结写入 Sheet1 的前几行
                    summary_df = pd.DataFrame([["DeepSeek R1 Summary"], [self.cached_summary], [""]])
                    summary_df.to_excel(writer, sheet_name='Report', index=False, header=False, startrow=0)
                    
                    # 将结果表格写入总结下方
                    pd.DataFrame([["Top 10 Reranked Results"]]).to_excel(writer, sheet_name='Report', index=False, header=False, startrow=4)
                    df.to_excel(writer, sheet_name='Report', index=False, startrow=6)
            
            # 2. 导出 CSV
            elif fmt == "csv":
                with open(save_path, 'w', newline='', encoding='utf-8-sig') as f:
                    writer = csv.writer(f)
                    # 写入总结
                    writer.writerow(["=== DeepSeek R1 Summary ==="])
                    writer.writerow([self.cached_summary])
                    writer.writerow([])
                    writer.writerow(["=== Top 10 Reranked Results ==="])
                    # 写入表头
                    writer.writerow(["Rank", "Final Score", "Debug Score", "Section Path", "Content"])
                    # 写入数据
                    for item in self.cached_results:
                        writer.writerow([
                            item['rank'],
                            f"{item['final_score']:.4f}",
                            item.get('debug_score', ''),
                            item['path'],
                            item['content']
                        ])

            # 3. 导出 TXT
            elif fmt == "txt":
                with open(save_path, 'w', encoding='utf-8') as f:
                    f.write(f"Query: {self.cached_query}\n")
                    f.write("="*50 + "\n")
                    f.write("DeepSeek R1 Summary:\n")
                    f.write("="*50 + "\n")
                    # 去除 Markdown 符号简化文本
                    clean_summary = self.cached_summary.replace("**", "").replace(">", "")
                    f.write(clean_summary + "\n\n")
                    
                    f.write("="*50 + "\n")
                    f.write("Top 10 Reranked Results:\n")
                    f.write("="*50 + "\n")
                    for item in self.cached_results:
                        f.write(f"[Rank #{item['rank']}] Score: {item['final_score']:.4f}\n")
                        f.write(f"Path: {item['path']}\n")
                        f.write(f"Content:\n{item['content']}\n")
                        f.write("-" * 30 + "\n")

            # 4. 导出 Markdown (MD)
            elif fmt == "md":
                with open(save_path, 'w', encoding='utf-8') as f:
                    f.write(f"# RAG Query Report\n\n")
                    f.write(f"**Query:** {self.cached_query}\n\n")
                    f.write(f"## 🤖 DeepSeek R1 Summary\n\n")
                    f.write(self.cached_summary + "\n\n")
                    f.write(f"## 📚 Top 10 Reranked Results\n\n")
                    for item in self.cached_results:
                        f.write(f"### Rank #{item['rank']} (Score: {item['final_score']:.4f})\n")
                        f.write(f"**Path:** `{item['path']}`\n\n")
                        f.write(f"**Content:**\n\n")
                        # 引用内容
                        content_block = item['content'].replace('\n', '\n> ')
                        f.write(f"> {content_block}\n\n")
                        f.write("---\n")

            # 5. 导出 DOCX
            elif fmt == "docx":
                if docx is None:
                    raise ImportError("缺少 python-docx 库，请 pip install python-docx")
                
                doc = docx.Document()
                doc.add_heading('RAG Analysis Report', 0)
                
                # Query
                p = doc.add_paragraph()
                p.add_run('Query: ').bold = True
                p.add_run(self.cached_query)
                
                # Summary
                doc.add_heading('DeepSeek R1 Summary', level=1)
                # 简单处理 Markdown 加粗逻辑用于 docx 展示
                # 这里简单写入文本，若需完美渲染 Markdown 需复杂解析
                doc.add_paragraph(self.cached_summary)
                
                # Results
                doc.add_heading('Top 10 Reranked Results', level=1)
                
                for item in self.cached_results:
                    p_header = doc.add_paragraph()
                    run = p_header.add_run(f"Rank #{item['rank']} | Score: {item['final_score']:.4f}")
                    run.bold = True
                    run.font.color.rgb = docx.shared.RGBColor(0, 100, 0)
                    
                    p_path = doc.add_paragraph()
                    p_path.add_run("Path: ").bold = True
                    p_path.add_run(item['path']).italic = True
                    
                    p_content = doc.add_paragraph(item['content'])
                    p_content.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    doc.add_paragraph("-" * 40)

            self.log(f"✅ 导出成功: {save_path}")
            QMessageBox.information(self, "成功", f"文件已导出至:\n{save_path}")

        except ImportError as ie:
            self.log(f"❌ 导出失败 (库缺失): {str(ie)}")
            QMessageBox.critical(self, "错误", f"导出失败，缺少必要库:\n{str(ie)}")
        except Exception as e:
            self.log(f"❌ 导出异常: {str(e)}")
            import traceback
            self.log(traceback.format_exc())
            QMessageBox.critical(self, "错误", f"导出过程中发生错误:\n{str(e)}")

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = RAGRecallApp()
    window.show()
    sys.exit(app.exec_())